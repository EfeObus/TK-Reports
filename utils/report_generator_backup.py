import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


class ReportGenerator:
    """Generates professional PDF and DOCX reports following Canadian government standards"""
    
    def __init__(self, config):
        """Initialize with report configuration"""
        self.config = config
    
    def create_docx_report(self, filepath, data_summary, narrative, charts, metadata):
        """Create comprehensive DOCX report following Canadian government format"""
        doc = Document()
        
        # Cover Page
        self._add_cover_page_docx(doc, metadata)
        
        # Executive Summary
        doc.add_page_break()
        heading = doc.add_heading('Executive Summary', level=1)
        if narrative and 'executive_summary' in narrative:
            doc.add_paragraph(narrative['executive_summary'])
        
        # Table of Contents
        doc.add_page_break()
        doc.add_heading('Table of Contents', level=1)
        toc_items = [
            "1. Executive Summary",
            "2. Introduction",
            "   2.1 Background",
            "   2.2 Objectives and Scope",
            "   2.3 Data Sources",
            "3. Methodology",
            "   3.1 Data Collection and Preparation",
            "   3.2 Analytical Approach",
            "   3.3 Tools and Frameworks",
            "4. Findings and Results",
            "   4.1 Dataset Overview",
            "   4.2 Variables Analyzed",
            "   4.3 Descriptive Statistics",
            "   4.4 Visual Analysis",
            "5. Insights and Interpretation",
            "   5.1 Key Findings",
            "   5.2 Data Interpretation",
            "6. Recommendations",
            "7. Conclusion",
            "Appendices",
            "References"
        ]
        for item in toc_items:
            doc.add_paragraph(item, style='List Number' if item[0].isdigit() else 'Normal')
        
        # List of Figures
        doc.add_page_break()
        doc.add_heading('List of Figures and Tables', level=1)
        if charts:
            for idx, chart in enumerate(charts, 1):
                doc.add_paragraph(f"Figure {idx}: {chart['title']}", style='List Number')
        
        # Introduction (Enhanced - 3 detailed sections)
        doc.add_page_break()
        doc.add_heading('1. Introduction', level=1)
        
        # Calculate some key metrics for introduction
        numeric_cols = data_summary.get('numeric_column_names', [])
        categorical_cols = data_summary.get('categorical_column_names', [])
        total_rows = data_summary.get('total_rows', 0)
        total_cols = data_summary.get('total_columns', 0)
        
        doc.add_paragraph(
            f"The purpose of this report is to present a structured analytical insight into the dataset "
            f"comprising {total_rows:,} records across {total_cols} distinct variables. This comprehensive "
            f"analysis examines demographic, quantitative, and categorical dimensions of the data to identify "
            f"patterns, trends, and actionable intelligence that can support strategic decision-making."
        )
        
        doc.add_paragraph(
            f"The dataset reflects a rich information landscape with {len(numeric_cols)} numerical variables "
            f"providing quantitative measurements, and {len(categorical_cols)} categorical variables offering "
            f"qualitative dimensions for segmentation and classification analysis. By examining these factors "
            f"through multiple analytical lenses, this report aims to uncover both obvious and hidden insights "
            f"that may inform operational improvements, strategic planning, and risk mitigation."
        )
        
        # Add quality and scope context
        quality_score = data_summary.get('data_quality_score', 0)
        quality_desc = "excellent" if quality_score >= 90 else "good" if quality_score >= 75 else "acceptable"
        
        doc.add_paragraph(
            f"The analytical scope encompasses descriptive statistical analysis, correlation studies, "
            f"distribution profiling, and pattern recognition. With a data quality score of {quality_score:.1f}/100 "
            f"indicating {quality_desc} data integrity, the findings presented herein are based on a solid "
            f"foundation suitable for evidence-based decision support. This report follows Canadian government "
            f"reporting standards to ensure clarity, professionalism, and accessibility for stakeholders at all levels."
        )
        
        # Methodology (Enhanced with detailed sections)
        doc.add_page_break()
        doc.add_heading('2. Methodology', level=1)
        
        # Get recommended analysis type
        analysis_type = data_summary.get('recommended_analysis_type', 'Exploratory Data Analysis')
        
        doc.add_heading('2.1 Data Source and Scope', level=2)
        doc.add_paragraph(
            f"Data Source: Dataset comprising {total_rows:,} records and {total_cols} variables"
        )
        doc.add_paragraph(
            f"Scope: Comprehensive analysis covering all {len(numeric_cols)} numerical variables and "
            f"{len(categorical_cols)} categorical variables across the entire dataset"
        )
        doc.add_paragraph(
            f"Time Frame: Current snapshot analysis performed on {datetime.now().strftime('%B %d, %Y')}"
        )
        
        doc.add_heading('2.2 Analytical Tools and Frameworks', level=2)
        doc.add_paragraph("The analysis employed a robust technical stack:")
        doc.add_paragraph("Python 3.11+ - Core data processing and statistical computation", style='List Bullet')
        doc.add_paragraph("Pandas & NumPy - Data manipulation and numerical analysis", style='List Bullet')
        doc.add_paragraph("SciPy - Advanced statistical functions (skewness, kurtosis, outlier detection)", style='List Bullet')
        doc.add_paragraph("Matplotlib & Seaborn - Professional data visualization", style='List Bullet')
        doc.add_paragraph("OpenAI GPT-4 Turbo - AI-powered narrative generation and insight interpretation", style='List Bullet')
        doc.add_paragraph("TK Reports Analytics Engine - Automated pattern detection and recommendation generation", style='List Bullet')
        
        doc.add_heading('2.3 Analytical Approach', level=2)
        doc.add_paragraph(
            f"Based on the dataset characteristics, the recommended analytical methodology is: {analysis_type}. "
            f"This approach combines multiple analytical techniques:"
        )
        
        doc.add_paragraph("Descriptive Statistics: Calculation of mean, median, standard deviation, quartiles, and range for all numerical variables", style='List Number')
        doc.add_paragraph("Distribution Analysis: Skewness and kurtosis assessment to understand data shape and tail behavior", style='List Number')
        doc.add_paragraph("Outlier Detection: IQR-based identification of anomalous values for data quality assessment", style='List Number')
        doc.add_paragraph("Correlation Analysis: Pearson correlation coefficients computed for all numeric variable pairs", style='List Number')
        doc.add_paragraph("Categorical Profiling: Frequency analysis, diversity scoring, and concentration metrics", style='List Number')
        doc.add_paragraph("Visual Analytics: Multi-dimensional visualization including distributions, correlations, and relationships", style='List Number')
        doc.add_paragraph("AI-Powered Interpretation: Natural language generation for insights, patterns, and recommendations", style='List Number')
        
        doc.add_heading('2.4 Quality Assurance', level=2)
        missing_pct = data_summary.get('missing_data_summary', {}).get('missing_percentage', 0)
        doc.add_paragraph(
            f"Data quality assessment revealed {missing_pct:.2f}% missing values across the dataset. "
            f"Quality assurance procedures included validation of data types, identification of duplicate records, "
            f"and verification of value ranges for logical consistency. The overall data quality score of "
            f"{quality_score:.1f}/100 indicates {quality_desc} data integrity suitable for analytical purposes."
        )
        
        # Findings (Enhanced with comprehensive tables and analysis)
        doc.add_page_break()
        doc.add_heading('3. Key Findings', level=1)
        
        # Use AI-generated findings if available
        if narrative and 'key_findings' in narrative:
            doc.add_paragraph(narrative['key_findings'])
            doc.add_paragraph()  # Add spacing
        
        # 3.1 Dataset Summary Table
        doc.add_heading('3.1 Dataset Composition', level=2)
        
        comp_table = doc.add_table(rows=7, cols=2)
        comp_table.style = 'Light Grid Accent 1'
        
        comp_table.rows[0].cells[0].text = 'Metric'
        comp_table.rows[0].cells[1].text = 'Value'
        for cell in comp_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        comp_table.rows[1].cells[0].text = 'Total Records'
        comp_table.rows[1].cells[1].text = f"{total_rows:,}"
        comp_table.rows[2].cells[0].text = 'Total Variables'
        comp_table.rows[2].cells[1].text = f"{total_cols}"
        comp_table.rows[3].cells[0].text = 'Numeric Variables'
        comp_table.rows[3].cells[1].text = f"{len(numeric_cols)}"
        comp_table.rows[4].cells[0].text = 'Categorical Variables'
        comp_table.rows[4].cells[1].text = f"{len(categorical_cols)}"
        comp_table.rows[5].cells[0].text = 'Data Quality Score'
        comp_table.rows[5].cells[1].text = f"{quality_score:.1f}/100"
        comp_table.rows[6].cells[0].text = 'Missing Data Percentage'
        comp_table.rows[6].cells[1].text = f"{missing_pct:.2f}%"
        
        doc.add_paragraph()
        
        # 3.2 Key Numerical Statistics
        if data_summary.get('numeric_summary'):
            doc.add_heading('3.2 Key Numerical Statistics', level=2)
            
            # Create detailed stats table for top 5 numeric variables
            num_vars = list(data_summary['numeric_summary'].items())[:5]
            if num_vars:
                stats_table = doc.add_table(rows=len(num_vars) + 1, cols=6)
                stats_table.style = 'Light Grid Accent 1'
                
                # Header row
                headers = ['Variable', 'Mean', 'Median', 'Std Dev', 'Min', 'Max']
                for idx, header in enumerate(headers):
                    cell = stats_table.rows[0].cells[idx]
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Data rows
                for row_idx, (var_name, stats) in enumerate(num_vars, start=1):
                    stats_table.rows[row_idx].cells[0].text = var_name
                    stats_table.rows[row_idx].cells[1].text = f"{stats.get('mean', 0):.2f}"
                    stats_table.rows[row_idx].cells[2].text = f"{stats.get('median', 0):.2f}"
                    stats_table.rows[row_idx].cells[3].text = f"{stats.get('std', 0):.2f}"
                    stats_table.rows[row_idx].cells[4].text = f"{stats.get('min', 0):.2f}"
                    stats_table.rows[row_idx].cells[5].text = f"{stats.get('max', 0):.2f}"
                
                doc.add_paragraph()
                
                # Add interpretation paragraph
                doc.add_paragraph(
                    "The table above presents descriptive statistics for the primary numerical variables in the dataset. "
                    "These metrics provide insight into central tendency (mean, median), dispersion (standard deviation), "
                    "and the range of values observed in the data."
                )
                doc.add_paragraph()
        
        # 3.3 Categorical Variable Breakdown
        if data_summary.get('categorical_summary'):
            doc.add_heading('3.3 Categorical Variable Distribution', level=2)
            
            cat_vars = list(data_summary['categorical_summary'].items())[:5]
            for var_name, stats in cat_vars:
                doc.add_paragraph(f"{var_name}:", style='Heading 3')
                doc.add_paragraph(f"Unique Values: {stats.get('unique_values', 0)}")
                
                top_values = stats.get('top_values', {})
                if top_values:
                    doc.add_paragraph("Top Categories:")
                    total_count = sum(top_values.values())
                    for value, count in list(top_values.items())[:5]:
                        percentage = (count / total_count * 100) if total_count > 0 else 0
                        doc.add_paragraph(f"  • {value}: {count:,} ({percentage:.1f}%)", style='List Bullet')
                
                # Add concentration insight
                top_pct = stats.get('top_value_percentage', 0)
                if top_pct > 50:
                    top_category = list(top_values.keys())[0] if top_values else 'N/A'
                    doc.add_paragraph(
                        f"Note: '{top_category}' dominates this variable at {top_pct:.1f}% of all records.",
                        style='Intense Quote'
                    )
                doc.add_paragraph()
        
        # 3.4 Correlation Analysis
        if data_summary.get('correlations'):
            doc.add_heading('3.4 Correlation Patterns', level=2)
            
            corr_data = data_summary['correlations'][:8]
            if corr_data:
                corr_table = doc.add_table(rows=len(corr_data) + 1, cols=4)
                corr_table.style = 'Light Grid Accent 1'
                
                # Headers
                headers = ['Variable 1', 'Variable 2', 'Correlation', 'Strength']
                for idx, header in enumerate(headers):
                    cell = corr_table.rows[0].cells[idx]
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Data
                for row_idx, corr in enumerate(corr_data, start=1):
                    corr_table.rows[row_idx].cells[0].text = corr['var1']
                    corr_table.rows[row_idx].cells[1].text = corr['var2']
                    corr_table.rows[row_idx].cells[2].text = f"{corr['correlation']:.3f}"
                    corr_table.rows[row_idx].cells[3].text = corr.get('strength', 'N/A')
                
                doc.add_paragraph()
                doc.add_paragraph(
                    "Correlation coefficients range from -1 to +1, where values closer to ±1 indicate stronger "
                    "linear relationships. Positive correlations suggest variables move in the same direction, "
                    "while negative correlations indicate inverse relationships."
                )
                doc.add_paragraph()
        
        # 3.5 Visual Analysis and Chart Interpretation
        doc.add_heading('3.5 Visual Analysis and Interpretation', level=2)
        
        if charts and self.config.get('include_charts', True):
            for idx, chart in enumerate(charts, 1):
                if os.path.exists(chart['path']):
                    # Add chart
                    doc.add_paragraph(f"Figure {idx}: {chart['title']}", style='Caption')
                    doc.add_picture(chart['path'], width=Inches(6.0))
                    doc.add_paragraph()
                    
                    # Add detailed interpretation based on chart type
                    interpretation = chart.get('interpretation', '')
                    chart_title_lower = chart['title'].lower()
                    
                    if 'distribution' in chart_title_lower:
                        doc.add_paragraph(
                            f"Chart Interpretation: {interpretation} "
                            "The distribution chart reveals the spread and frequency of values within the variable. "
                            "The mean (red dashed line) indicates the average value, while the median (green dashed line) "
                            "shows the middle point of the data. When mean and median differ significantly, it suggests "
                            "the presence of skewness in the distribution. The shape of the distribution helps identify "
                            "whether data follows a normal pattern or exhibits unusual clustering or outliers."
                        )
                    
                    elif 'correlation' in chart_title_lower or 'heatmap' in chart_title_lower:
                        doc.add_paragraph(
                            f"Chart Interpretation: {interpretation} "
                            "The correlation matrix presents the strength and direction of linear relationships between "
                            "numerical variables. Darker colors (closer to 1 or -1) indicate strong correlations, while "
                            "lighter colors (closer to 0) suggest weak or no linear relationship. Positive correlations "
                            "(warmer colors) indicate variables that increase together, while negative correlations "
                            "(cooler colors) show inverse relationships. Identifying strong correlations is crucial for "
                            "understanding variable dependencies and potential multicollinearity in predictive models."
                        )
                    
                    elif 'category' in chart_title_lower or 'bar' in chart_title_lower:
                        doc.add_paragraph(
                            f"Chart Interpretation: {interpretation} "
                            "The categorical distribution chart displays the frequency and percentage of each category "
                            "within the variable. The height of bars represents the count of observations in each category, "
                            "while percentage labels provide relative proportions. This visualization helps identify dominant "
                            "categories, balanced vs. imbalanced distributions, and potential data collection biases. "
                            "Categories with very low frequencies may warrant consolidation or special treatment in analysis."
                        )
                    
                    elif 'box plot' in chart_title_lower or 'boxplot' in chart_title_lower:
                        doc.add_paragraph(
                            f"Chart Interpretation: {interpretation} "
                            "The box plot provides a comprehensive view of data distribution through quartiles. The box "
                            "represents the interquartile range (IQR) containing the middle 50% of data, with the line "
                            "inside marking the median. Whiskers extend to show the range of typical values, while points "
                            "beyond whiskers indicate potential outliers. This visualization is particularly effective for "
                            "comparing distributions across different categories and identifying data quality issues."
                        )
                    
                    elif 'scatter' in chart_title_lower:
                        doc.add_paragraph(
                            f"Chart Interpretation: {interpretation} "
                            "The scatter plot illustrates the relationship between two numerical variables, with each point "
                            "representing an individual observation. The trend line (if present) shows the general direction "
                            "and strength of the relationship. Patterns in the scatter can reveal linear relationships, "
                            "non-linear associations, or clusters of similar observations. The spread of points around the "
                            "trend line indicates variability and the strength of the relationship between variables."
                        )
                    
                    else:
                        # Generic interpretation for other chart types
                        doc.add_paragraph(f"Chart Interpretation: {interpretation}")
                    
                    doc.add_paragraph()  # Add spacing between charts
        else:
            doc.add_paragraph(
                "Visual analysis charts were generated to illustrate key patterns and relationships in the data. "
                "Each chart provides unique insights into different aspects of the dataset."
            )
            doc.add_paragraph()

        
        # Insights
        doc.add_page_break()
        doc.add_heading('4. Insights and Interpretation', level=1)
        
        if narrative and 'insights_interpretation' in narrative:
            # Use the comprehensive AI-generated insights
            doc.add_paragraph(narrative['insights_interpretation'])
        else:
            # Fallback to key findings
            doc.add_heading('4.1 Key Findings', level=2)
            if narrative and 'key_findings' in narrative:
                doc.add_paragraph(narrative['key_findings'])
            
            doc.add_heading('4.2 Data Interpretation', level=2)
            doc.add_paragraph(
                "The analysis reveals significant patterns and relationships within the dataset. "
                "These findings provide a foundation for evidence-based decision making."
            )
        
        # Recommendations
        doc.add_page_break()
        doc.add_heading('5. Recommendations', level=1)
        if narrative and 'recommendations' in narrative:
            doc.add_paragraph(narrative['recommendations'])
        else:
            doc.add_paragraph(
                "Based on the analysis conducted, the following recommendations are proposed to leverage "
                "insights from the data and address identified patterns and opportunities."
            )
        
        # Conclusion (Enhanced with comprehensive structure)
        doc.add_page_break()
        doc.add_heading('6. Conclusion', level=1)
        
        # 6.1 Summary of Key Insights
        doc.add_heading('6.1 Summary of Key Insights', level=2)
        doc.add_paragraph(
            f"This comprehensive analysis examined a dataset containing {total_rows:,} records across "
            f"{total_cols} variables, employing advanced statistical techniques and AI-powered interpretation. "
            f"The study revealed significant patterns and relationships that provide actionable intelligence "
            "for strategic decision-making."
        )
        doc.add_paragraph()
        
        # Summarize key findings
        key_insights = []
        
        if data_summary.get('numeric_summary'):
            num_count = len(data_summary['numeric_summary'])
            key_insights.append(
                f"Analyzed {num_count} numerical variables using descriptive statistics, distribution analysis, "
                "and outlier detection to understand central tendencies, variability, and data quality."
            )
        
        if data_summary.get('categorical_summary'):
            cat_count = len(data_summary['categorical_summary'])
            key_insights.append(
                f"Examined {cat_count} categorical variables to identify dominant categories, distribution patterns, "
                "and potential imbalances that may impact analysis outcomes."
            )
        
        if data_summary.get('correlations'):
            strong_corrs = [c for c in data_summary['correlations'] if abs(c['correlation']) > 0.7]
            if strong_corrs:
                key_insights.append(
                    f"Identified {len(strong_corrs)} strong correlations (|r| > 0.7) between variables, "
                    "indicating significant linear relationships that warrant further investigation."
                )
        
        if quality_score >= 90:
            key_insights.append(
                f"The dataset demonstrated excellent data quality (score: {quality_score:.1f}/100), "
                "providing a reliable foundation for analysis and decision-making."
            )
        elif quality_score >= 70:
            key_insights.append(
                f"The dataset showed good data quality (score: {quality_score:.1f}/100) with minor gaps "
                "that were addressed through appropriate handling of missing values."
            )
        else:
            key_insights.append(
                f"Data quality analysis (score: {quality_score:.1f}/100) revealed areas for improvement, "
                "particularly regarding data completeness and consistency."
            )
        
        for insight in key_insights:
            doc.add_paragraph(insight, style='List Bullet')
        
        doc.add_paragraph()
        
        # 6.2 Implications and Strategic Value
        doc.add_heading('6.2 Implications and Strategic Value', level=2)
        doc.add_paragraph(
            "The findings from this analysis carry significant implications for organizational strategy and operations. "
            "The identified patterns and relationships provide evidence-based insights that can inform policy development, "
            "resource allocation, and operational improvements. The comprehensive statistical foundation established through "
            "this study enables stakeholders to make informed decisions with confidence in the underlying data quality and "
            "analytical rigor."
        )
        doc.add_paragraph()
        doc.add_paragraph(
            "Moreover, the correlation analysis reveals interdependencies between variables that may have cascading effects "
            "across different operational domains. Understanding these relationships is crucial for developing holistic "
            "strategies that account for systemic impacts rather than isolated interventions."
        )
        doc.add_paragraph()
        
        # 6.3 Limitations and Considerations
        doc.add_heading('6.3 Limitations and Considerations', level=2)
        doc.add_paragraph(
            "While this analysis provides comprehensive insights, several limitations should be acknowledged:"
        )
        
        limitations = [
            "The analysis is based on historical data and assumes that observed patterns will remain relevant. "
            "Changes in underlying conditions or external factors may affect the applicability of findings.",
            
            "Statistical correlations identified in this study indicate associations but do not establish causation. "
            "Additional research may be needed to understand causal mechanisms.",
            
            f"Data quality assessment revealed a quality score of {quality_score:.1f}/100 and {missing_pct:.2f}% missing data. "
            "While these metrics are within acceptable ranges, missing data handling methods may introduce minor biases.",
            
            "The analysis focuses on linear relationships and standard statistical measures. Non-linear patterns or "
            "complex interactions may require specialized analytical techniques.",
            
            "Interpretation of findings should consider the specific context of data collection, including timing, "
            "methodology, and any sampling limitations that may affect generalizability."
        ]
        
        for limitation in limitations:
            doc.add_paragraph(limitation, style='List Bullet')
        
        doc.add_paragraph()
        
        # 6.4 Next Steps and Future Research
        doc.add_heading('6.4 Next Steps and Future Research', level=2)
        doc.add_paragraph(
            "Building on the foundation established by this analysis, the following next steps are recommended:"
        )
        
        next_steps = [
            "Validate findings through stakeholder consultation and subject matter expert review to ensure "
            "alignment with organizational objectives and operational realities.",
            
            "Develop targeted action plans based on specific recommendations, with clear ownership, timelines, "
            "and success metrics.",
            
            "Establish monitoring frameworks to track key variables and trends over time, enabling early detection "
            "of changes and assessment of intervention effectiveness.",
            
            "Conduct deeper analysis of identified correlations and patterns to understand underlying mechanisms "
            "and develop predictive models where appropriate.",
            
            "Consider integrating additional data sources to enrich the analytical foundation and enable "
            "more comprehensive insights."
        ]
        
        for step in next_steps:
            doc.add_paragraph(step, style='List Bullet')
        
        doc.add_paragraph()
        
        # Closing statement
        doc.add_paragraph(
            "This report represents a rigorous, evidence-based examination of the dataset using industry-standard "
            "analytical techniques and AI-powered interpretation. The insights and recommendations presented herein "
            "are designed to support informed decision-making and drive positive organizational outcomes. "
            "The TK Reports Engine ensures that analysis adheres to Canadian government standards for professional "
            "reporting, providing stakeholders with reliable, actionable intelligence."
        )

        
        # Appendices
        doc.add_page_break()
        doc.add_heading('Appendices', level=1)
        
        doc.add_heading('Appendix A: Statistical Correlations', level=2)
        if data_summary.get('correlations'):
            for corr in data_summary['correlations'][:10]:
                doc.add_paragraph(
                    f"{corr['var1']} ↔ {corr['var2']}: {corr['correlation']:.3f}",
                    style='List Bullet'
                )
        
        doc.add_heading('Appendix B: Categorical Summaries', level=2)
        if data_summary.get('categorical_summary'):
            for col, stats in list(data_summary['categorical_summary'].items())[:5]:
                doc.add_paragraph(f"{col}:", style='List Bullet')
                top_values = stats.get('top_values', {})
                for value, count in list(top_values.items())[:3]:
                    doc.add_paragraph(f"  {value}: {count}")
        
        # References
        doc.add_page_break()
        doc.add_heading('References', level=1)
        doc.add_paragraph("This section can be populated with relevant references following APA 7th edition format.")
        
        # Save document
        output_path = filepath.replace('.pdf', '.docx') if filepath.endswith('.pdf') else filepath
        doc.save(output_path)
        print(f"DOCX report saved to: {output_path}")
    
    def _add_cover_page_docx(self, doc, metadata):
        """Add professional cover page to DOCX"""
        title = doc.add_heading(metadata.get('title', 'Data Insight Report'), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('AI-Powered Data Analysis Report', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        meta_table = doc.add_table(rows=4, cols=2)
        meta_table.style = 'Light Grid Accent 1'
        meta_table.rows[0].cells[0].text = 'Prepared For:'
        meta_table.rows[0].cells[1].text = metadata.get('prepared_for', 'N/A')
        meta_table.rows[1].cells[0].text = 'Prepared By:'
        meta_table.rows[1].cells[1].text = metadata.get('prepared_by', 'TK Reports')
        meta_table.rows[2].cells[0].text = 'Date of Submission:'
        meta_table.rows[2].cells[1].text = metadata.get('date', datetime.now().strftime('%Y-%m-%d'))
        meta_table.rows[3].cells[0].text = 'Locale:'
        meta_table.rows[3].cells[1].text = metadata.get('locale', 'en_CA')
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        notice = doc.add_paragraph()
        notice_run = notice.add_run('CONFIDENTIALITY NOTICE')
        notice_run.font.bold = True
        notice_run.font.size = Pt(10)
        notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        notice_text = doc.add_paragraph(
            'This document contains confidential information intended solely for the use of the individual '
            'or entity named above. Any unauthorized review, use, disclosure, or distribution is prohibited.'
        )
        notice_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def create_pdf_report(self, filepath, data_summary, narrative, charts, metadata):
        """Create PDF report following Canadian government format"""
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Cover Page
        pdf.add_page()
        self._add_cover_page_pdf(pdf, metadata)
        
        # Executive Summary
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, 'Executive Summary', ln=True)
        pdf.ln(5)
        pdf.set_font('Arial', '', 11)
        if narrative and 'executive_summary' in narrative:
            clean_text = self._clean_text_for_pdf(narrative['executive_summary'])
            pdf.multi_cell(0, 6, clean_text)
        
        # Table of Contents
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, 'Table of Contents', ln=True)
        pdf.ln(5)
        pdf.set_font('Arial', '', 11)
        for item in ["1. Executive Summary", "2. Introduction", "3. Methodology", "4. Findings", "5. Insights", "6. Recommendations", "7. Conclusion"]:
            pdf.cell(0, 8, item, ln=True)
        
        # Introduction
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, '1. Introduction', ln=True)
        pdf.ln(5)
        pdf.set_font('Arial', '', 11)
        intro_text = (
            f"This report presents comprehensive analysis of the provided dataset containing "
            f"{data_summary.get('total_rows', 'N/A'):,} records and {data_summary.get('total_columns', 'N/A')} variables."
        )
        pdf.multi_cell(0, 6, intro_text)
        
        # Findings
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, '3. Findings and Results', ln=True)
        pdf.ln(5)
        
        pdf.set_font('Arial', 'B', 11)
        pdf.cell(70, 8, 'Total Records:', border=1)
        pdf.set_font('Arial', '', 11)
        pdf.cell(0, 8, f"{data_summary.get('total_rows', 'N/A'):,}", border=1, ln=True)
        
        # Charts
        if charts and self.config.get('include_charts', True):
            for idx, chart in enumerate(charts, 1):
                pdf.add_page()
                pdf.set_font('Arial', 'B', 14)
                pdf.cell(0, 10, f'Figure {idx}: {chart["title"]}', ln=True)
                pdf.ln(5)
                if os.path.exists(chart['path']):
                    pdf.image(chart['path'], x=10, w=190)
        
        # Insights
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, '4. Insights and Interpretation', ln=True)
        pdf.ln(5)
        pdf.set_font('Arial', '', 11)
        
        if narrative and 'insights_interpretation' in narrative:
            clean_text = self._clean_text_for_pdf(narrative['insights_interpretation'])
            pdf.multi_cell(0, 6, clean_text)
        elif narrative and 'key_findings' in narrative:
            clean_text = self._clean_text_for_pdf(narrative['key_findings'])
            pdf.multi_cell(0, 6, clean_text)
        
        # Recommendations
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, '5. Recommendations', ln=True)
        pdf.ln(5)
        if narrative and 'recommendations' in narrative:
            pdf.set_font('Arial', '', 11)
            clean_text = self._clean_text_for_pdf(narrative['recommendations'])
            pdf.multi_cell(0, 6, clean_text)
        
        output_path = filepath if filepath.endswith('.pdf') else filepath.replace('.docx', '.pdf')
        pdf.output(output_path)
        print(f"PDF report saved to: {output_path}")
    
    def _add_cover_page_pdf(self, pdf, metadata):
        """Add professional cover page to PDF"""
        pdf.set_font('Arial', 'B', 24)
        pdf.ln(40)
        pdf.cell(0, 20, metadata.get('title', 'Data Insight Report'), ln=True, align='C')
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, 'AI-Powered Data Analysis Report', ln=True, align='C')
        pdf.ln(30)
        
        pdf.set_font('Arial', '', 12)
        pdf.cell(70, 10, 'Prepared For:', border=1)
        pdf.cell(0, 10, metadata.get('prepared_for', 'N/A'), border=1, ln=True)
        pdf.cell(70, 10, 'Prepared By:', border=1)
        pdf.cell(0, 10, metadata.get('prepared_by', 'TK Reports'), border=1, ln=True)
        pdf.cell(70, 10, 'Date:', border=1)
        pdf.cell(0, 10, metadata.get('date', datetime.now().strftime('%Y-%m-%d')), border=1, ln=True)
        
        pdf.ln(30)
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(0, 10, 'CONFIDENTIALITY NOTICE', ln=True, align='C')
        pdf.set_font('Arial', '', 9)
        pdf.multi_cell(0, 5, 'This document contains confidential information.', align='C')
    
    def _clean_text_for_pdf(self, text):
        """Clean text to remove problematic Unicode characters for PDF"""
        if not text:
            return ""
        
        # Replace common problematic characters
        replacements = {
            '•': '-',  # Bullet point to dash
            '–': '-',  # En dash to hyphen
            '—': '-',  # Em dash to hyphen
            '"': '"',  # Smart quotes to regular
            '"': '"',
            ''': "'",
            ''': "'",
            '…': '...',  # Ellipsis
            '→': '->',  # Arrow
            '↔': '<->',  # Double arrow
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        return text
