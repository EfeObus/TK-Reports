import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


class ReportGenerator:
    """Generates comprehensive professional reports following TK Reports standards"""
    
    def __init__(self, config):
        """Initialize with report configuration"""
        self.config = config
    
    def create_docx_report(self, filepath, data_summary, narrative, charts, metadata):
        """Create comprehensive DOCX report with detailed findings and interpretations"""
        doc = Document()
        
        # Cover Page
        self._add_cover_page_docx(doc, metadata)
        
        # Extract key metrics for use throughout report
        total_rows = data_summary.get('total_rows', 0)
        total_cols = data_summary.get('total_columns', 0)
        numeric_cols = data_summary.get('numeric_column_names', [])
        categorical_cols = data_summary.get('categorical_column_names', [])
        numeric_summary = data_summary.get('numeric_summary', {})
        categorical_summary = data_summary.get('categorical_summary', {})
        correlations = data_summary.get('correlations', [])
        quality_score = data_summary.get('data_quality_score', 0)
        missing_pct = data_summary.get('missing_percentage', 0)
        insights = data_summary.get('insights', [])
        
        # 1. EXECUTIVE SUMMARY
        doc.add_page_break()
        doc.add_heading('1. Executive Summary', level=1)
        
        if narrative and 'executive_summary' in narrative:
            doc.add_paragraph(narrative['executive_summary'])
        else:
            # Generate comprehensive executive summary from data
            self._add_executive_summary(doc, data_summary, numeric_summary, categorical_summary, correlations)
        
        # 2. INTRODUCTION
        doc.add_page_break()
        doc.add_heading('2. Introduction', level=1)
        
        # Add Research Objective if provided
        if metadata.get('objective'):
            doc.add_heading('2.1 Research Objective and Purpose', level=2)
            doc.add_paragraph(metadata['objective'])
            doc.add_paragraph()
            doc.add_heading('2.2 Dataset Overview', level=2)
        
        doc.add_paragraph(
            f"This comprehensive analytical report examines a dataset containing {total_rows:,} records "
            f"across {total_cols} variables, encompassing both numerical metrics and categorical dimensions. "
            f"The dataset provides a multifaceted view with {len(numeric_cols)} quantitative variables "
            f"measuring continuous metrics, and {len(categorical_cols)} categorical variables capturing "
            f"qualitative attributes and classifications."
        )
        
        section_label = '2.3' if metadata.get('objective') else '2.1'
        doc.add_heading(f'{section_label} Analytical Objectives', level=2)
        doc.add_paragraph(
            "The analysis addresses the following key objectives:"
        )
        
        objectives = [
            "Understanding the composition and structure of the dataset to identify key characteristics and patterns.",
            "Examining relationships between variables to uncover correlations and dependencies that may inform strategic decisions.",
            "Identifying trends, anomalies, and distributions that reveal operational insights and risk factors.",
            "Providing evidence-based recommendations grounded in statistical analysis and industry best practices."
        ]
        
        for obj in objectives:
            doc.add_paragraph(obj, style='List Bullet')
        
        doc.add_paragraph()
        doc.add_paragraph(
            f"With a data quality score of {quality_score:.1f}/100 and {missing_pct:.2f}% missing values, "
            f"the dataset demonstrates {'excellent' if quality_score >= 90 else 'good' if quality_score >= 75 else 'acceptable'} "
            f"integrity, providing a reliable foundation for comprehensive analysis and actionable insights."
        )
        
        # 3. METHODOLOGY
        doc.add_page_break()
        doc.add_heading('3. Methodology', level=1)
        
        doc.add_heading('3.1 Data Source and Scope', level=2)
        doc.add_paragraph(
            f"The source dataset comprises {total_rows:,} individual records with {total_cols} distinct "
            f"variables. Analysis was conducted on {len(numeric_cols)} numerical fields and "
            f"{len(categorical_cols)} categorical dimensions. Data quality assessment revealed "
            f"{missing_pct:.2f}% missing values, which were appropriately handled through statistical "
            f"imputation or exclusion based on analytical context."
        )
        
        doc.add_heading('3.2 Analytical Tools and Technologies', level=2)
        tools = [
            "TK Reports Insight AI Engine — AI-assisted analytical platform with GPT-4 Turbo integration",
            "Python 3.11+ — Core analytical programming environment",
            "Pandas 2.x & NumPy — Data manipulation and numerical computations",
            "SciPy — Advanced statistical analysis (skewness, kurtosis, outlier detection)",
            "Matplotlib & Seaborn — Statistical visualization and graphical analysis",
            "Canadian Locale Standards (en_CA) — Formatting, currency, and date representations"
        ]
        for tool in tools:
            doc.add_paragraph(tool, style='List Bullet')
        
        doc.add_heading('3.3 Analytical Process', level=2)
        doc.add_paragraph(
            "The analysis followed a structured multi-stage approach:"
        )
        
        process_steps = [
            "Descriptive Statistics — Calculated mean, median, standard deviation, quartiles, skewness, "
            "and kurtosis for all numerical variables to understand central tendency and dispersion.",
            
            "Distribution Analysis — Examined frequency distributions, normality patterns, and outlier "
            "detection using Interquartile Range (IQR) methodology.",
            
            "Advanced Statistical Metrics — Computed coefficient of variation for relative variability "
            "assessment and Shannon entropy for categorical diversity measurement.",
            
            "Correlation Analysis — Applied Pearson correlation coefficients to identify linear "
            "relationships between numerical variables, with strength classification (Very Strong: |r| > 0.8, "
            "Strong: 0.6-0.8, Moderate: 0.4-0.6, Weak: 0.2-0.4, Very Weak: < 0.2).",
            
            "Categorical Profiling — Analyzed frequency distributions, concentration metrics, and dominant "
            "category identification across all categorical variables.",
            
            "Visual Analytics — Generated five intelligent visualizations including distribution plots, "
            "correlation heatmaps, categorical bar charts, box plots for outlier detection, and scatter "
            "plots with trend lines.",
            
            "AI-Powered Interpretation — Leveraged GPT-4 Turbo to synthesize findings into comprehensive "
            "narratives, contextualizing statistical results within industry standards and best practices."
        ]
        
        for idx, step in enumerate(process_steps, 1):
            doc.add_paragraph(f"{idx}. {step}")
        
        doc.add_heading('3.4 Quality Assurance', level=2)
        doc.add_paragraph(
            f"Data validation procedures confirmed {quality_score:.1f}% overall data quality. Missing value "
            f"handling followed industry best practices, with systematic exclusion or imputation based on "
            f"variable criticality and missing data patterns. All statistical computations were validated "
            f"through cross-verification and outlier analysis to ensure accuracy and reliability."
        )
        
        # 4. DETAILED FINDINGS AND INTERPRETATIONS
        doc.add_page_break()
        doc.add_heading('4. Detailed Findings and Interpretations', level=1)
        
        # Use AI-generated key findings if available
        if narrative and 'key_findings' in narrative:
            doc.add_paragraph(narrative['key_findings'])
            doc.add_paragraph()
        
        # 4.1 Dataset Composition
        doc.add_heading('4.1 Dataset Composition and Overview', level=2)
        
        doc.add_paragraph(
            f"Observation: The dataset encompasses {total_rows:,} individual records representing the "
            f"analytical universe, with {total_cols} variables providing multidimensional perspectives."
        )
        
        doc.add_paragraph(
            f"Composition Insight: Of the {total_cols} variables, {len(numeric_cols)} ({len(numeric_cols)/total_cols*100:.1f}%) "
            f"are numerical metrics enabling quantitative analysis, while {len(categorical_cols)} "
            f"({len(categorical_cols)/total_cols*100:.1f}%) are categorical attributes supporting segmentation "
            f"and classification analysis."
        )
        
        doc.add_paragraph(
            f"Interpretation: This balanced composition of numerical and categorical variables enables "
            f"comprehensive multi-dimensional analysis. The {len(numeric_cols)}:{len(categorical_cols)} "
            f"ratio suggests a dataset well-suited for both statistical modeling and categorical profiling, "
            f"typical of {'operational' if len(numeric_cols) > len(categorical_cols) else 'descriptive'} "
            f"business intelligence datasets."
        )
        
        # Composition Summary Table
        doc.add_paragraph()
        comp_table = doc.add_table(rows=7, cols=2)
        comp_table.style = 'Light Grid Accent 1'
        
        comp_table.rows[0].cells[0].text = 'Metric'
        comp_table.rows[0].cells[1].text = 'Value'
        for cell in comp_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        comp_table.rows[1].cells[0].text = 'Total Records'
        comp_table.rows[1].cells[1].text = f"{total_rows:,}"
        comp_table.rows[2].cells[0].text = 'Total Variables'
        comp_table.rows[2].cells[1].text = f"{total_cols}"
        comp_table.rows[3].cells[0].text = 'Numerical Variables'
        comp_table.rows[3].cells[1].text = f"{len(numeric_cols)} ({len(numeric_cols)/total_cols*100:.1f}%)"
        comp_table.rows[4].cells[0].text = 'Categorical Variables'
        comp_table.rows[4].cells[1].text = f"{len(categorical_cols)} ({len(categorical_cols)/total_cols*100:.1f}%)"
        comp_table.rows[5].cells[0].text = 'Data Quality Score'
        comp_table.rows[5].cells[1].text = f"{quality_score:.1f}/100"
        comp_table.rows[6].cells[0].text = 'Missing Data'
        comp_table.rows[6].cells[1].text = f"{missing_pct:.2f}%"
        
        doc.add_paragraph()
        
        # 4.2 Numerical Variables - Detailed Analysis
        if numeric_summary:
            doc.add_heading('4.2 Key Numerical Variables — Statistical Profile', level=2)
            
            # Detailed stats table for all numeric variables (or top 10)
            num_vars = list(numeric_summary.items())[:10]
            
            stats_table = doc.add_table(rows=len(num_vars) + 1, cols=7)
            stats_table.style = 'Light Grid Accent 1'
            
            # Header row
            headers = ['Variable', 'Mean', 'Median', 'Std Dev', 'Min', 'Max', 'Range']
            for idx, header in enumerate(headers):
                cell = stats_table.rows[0].cells[idx]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Data rows
            for row_idx, (var_name, stats) in enumerate(num_vars, start=1):
                mean_val = stats.get('mean', 0)
                median_val = stats.get('median', 0)
                std_val = stats.get('std', 0)
                min_val = stats.get('min', 0)
                max_val = stats.get('max', 0)
                range_val = max_val - min_val
                
                stats_table.rows[row_idx].cells[0].text = var_name
                stats_table.rows[row_idx].cells[1].text = f"{mean_val:,.2f}"
                stats_table.rows[row_idx].cells[2].text = f"{median_val:,.2f}"
                stats_table.rows[row_idx].cells[3].text = f"{std_val:,.2f}"
                stats_table.rows[row_idx].cells[4].text = f"{min_val:,.2f}"
                stats_table.rows[row_idx].cells[5].text = f"{max_val:,.2f}"
                stats_table.rows[row_idx].cells[6].text = f"{range_val:,.2f}"
            
            doc.add_paragraph()
            
            # Add detailed interpretation for each major variable
            doc.add_paragraph(
                "Trend Insights: The statistical profile reveals distinct patterns across numerical variables:"
            )
            
            for var_name, stats in num_vars[:5]:  # Detailed analysis for top 5
                mean_val = stats.get('mean', 0)
                median_val = stats.get('median', 0)
                std_val = stats.get('std', 0)
                min_val = stats.get('min', 0)
                max_val = stats.get('max', 0)
                skew_val = stats.get('skewness', 0)
                
                # Calculate coefficient of variation
                cv = (std_val / mean_val * 100) if mean_val != 0 else 0
                
                # Mean vs Median comparison
                diff_pct = abs(mean_val - median_val) / mean_val * 100 if mean_val != 0 else 0
                
                doc.add_paragraph(f"• {var_name}:", style='List Bullet')
                
                interpretation = f"  Average value of {mean_val:,.2f} with median at {median_val:,.2f}"
                
                if diff_pct > 10:
                    interpretation += f" (divergence of {diff_pct:.1f}% indicates {'positive' if mean_val > median_val else 'negative'} skewness)"
                else:
                    interpretation += " (close alignment suggests symmetrical distribution)"
                
                interpretation += f". Standard deviation of {std_val:,.2f} yields coefficient of variation of {cv:.1f}%, "
                interpretation += f"indicating {'high' if cv > 50 else 'moderate' if cv > 20 else 'low'} relative variability. "
                interpretation += f"Value range spans {min_val:,.2f} to {max_val:,.2f}"
                
                if abs(skew_val) > 1:
                    interpretation += f", with skewness of {skew_val:.2f} suggesting {'right-tailed' if skew_val > 0 else 'left-tailed'} distribution"
                
                interpretation += "."
                
                doc.add_paragraph(interpretation)
            
            doc.add_paragraph()
        
        # 4.3 Categorical Variables - Distribution Analysis
        if categorical_summary:
            doc.add_heading('4.3 Categorical Variables — Distribution and Concentration', level=2)
            
            cat_vars = list(categorical_summary.items())[:5]
            
            for var_name, stats in cat_vars:
                unique_count = stats.get('unique_values', 0)
                top_values = stats.get('top_values', {})
                top_pct = stats.get('top_value_percentage', 0)
                
                doc.add_paragraph(f"{var_name}:", style='Heading 3')
                
                doc.add_paragraph(
                    f"Observation: This variable contains {unique_count} distinct categories across "
                    f"{total_rows:,} records, with the dominant category representing {top_pct:.1f}% "
                    f"of all observations."
                )
                
                if top_values:
                    doc.add_paragraph("Category Distribution:")
                    total_count = sum(top_values.values())
                    
                    # Create distribution table
                    dist_table = doc.add_table(rows=min(len(top_values), 5) + 1, cols=3)
                    dist_table.style = 'Light Grid Accent 1'
                    
                    dist_table.rows[0].cells[0].text = 'Category'
                    dist_table.rows[0].cells[1].text = 'Count'
                    dist_table.rows[0].cells[2].text = 'Percentage'
                    
                    for cell in dist_table.rows[0].cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    for idx, (value, count) in enumerate(list(top_values.items())[:5], start=1):
                        percentage = (count / total_count * 100) if total_count > 0 else 0
                        dist_table.rows[idx].cells[0].text = str(value)
                        dist_table.rows[idx].cells[1].text = f"{count:,}"
                        dist_table.rows[idx].cells[2].text = f"{percentage:.1f}%"
                
                doc.add_paragraph()
                
                # Interpretation based on distribution
                if top_pct > 70:
                    top_cat = list(top_values.keys())[0] if top_values else 'N/A'
                    doc.add_paragraph(
                        f"Interpretation: Extreme concentration with '{top_cat}' dominating at {top_pct:.1f}%. "
                        f"This high concentration (>70%) may indicate limited diversity in this dimension, "
                        f"potentially affecting segmentation analysis. Consider consolidating minor categories "
                        f"or investigating data collection biases.",
                        style='Intense Quote'
                    )
                elif top_pct > 50:
                    doc.add_paragraph(
                        f"Trend Insight: Moderate concentration at {top_pct:.1f}% suggests a dominant category "
                        f"with substantial but not overwhelming prevalence. This distribution is typical of "
                        f"organizational structures with clear hierarchies or market segments with leading players."
                    )
                elif unique_count <= 5:
                    doc.add_paragraph(
                        f"Insight: With only {unique_count} categories and balanced distribution ({top_pct:.1f}% maximum), "
                        f"this variable exhibits even representation suitable for comparative analysis and stratification."
                    )
                else:
                    doc.add_paragraph(
                        f"Insight: High diversity with {unique_count} categories and moderate concentration ({top_pct:.1f}%). "
                        f"This distribution pattern enables granular segmentation while maintaining analytical tractability."
                    )
                
                doc.add_paragraph()
        
        # 4.4 Correlation Analysis - Quantitative Relationships
        if correlations:
            doc.add_heading('4.4 Quantitative Correlations and Variable Relationships', level=2)
            
            doc.add_paragraph(
                "Observation: Correlation analysis reveals the strength and direction of linear relationships "
                "between numerical variables, identifying key dependencies and potential multicollinearity."
            )
            
            # Correlation table
            corr_data = correlations[:10]
            corr_table = doc.add_table(rows=len(corr_data) + 1, cols=4)
            corr_table.style = 'Light Grid Accent 1'
            
            # Headers
            headers = ['Variable Pair', 'Correlation (r)', 'Strength', 'Interpretation']
            for idx, header in enumerate(headers):
                cell = corr_table.rows[0].cells[idx]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Data
            for row_idx, corr in enumerate(corr_data, start=1):
                var_pair = f"{corr['var1']} ↔ {corr['var2']}"
                corr_val = corr['correlation']
                strength = corr.get('strength', 'N/A')
                
                # Interpretation
                if abs(corr_val) > 0.7:
                    interp = "Strong linear relationship" if corr_val > 0 else "Strong inverse relationship"
                elif abs(corr_val) > 0.4:
                    interp = "Moderate association" if corr_val > 0 else "Moderate negative association"
                else:
                    interp = "Weak relationship"
                
                corr_table.rows[row_idx].cells[0].text = var_pair
                corr_table.rows[row_idx].cells[1].text = f"{corr_val:.3f}"
                corr_table.rows[row_idx].cells[2].text = strength
                corr_table.rows[row_idx].cells[3].text = interp
            
            doc.add_paragraph()
            
            # Detailed interpretation of strongest correlations
            strong_corrs = [c for c in correlations if abs(c['correlation']) > 0.5]
            
            if strong_corrs:
                doc.add_paragraph(
                    f"Trend Insight: {len(strong_corrs)} correlation{'s' if len(strong_corrs) != 1 else ''} "
                    f"with |r| > 0.5 indicate{'s' if len(strong_corrs) == 1 else ''} substantial linear relationships:"
                )
                
                for corr in strong_corrs[:5]:  # Detail top 5
                    doc.add_paragraph(
                        f"• {corr['var1']} ↔ {corr['var2']} (r = {corr['correlation']:.3f}): "
                        f"{'Positive' if corr['correlation'] > 0 else 'Negative'} correlation suggests "
                        f"{'variables move in tandem' if corr['correlation'] > 0 else 'inverse relationship'}. "
                        f"This {corr.get('strength', 'significant').lower()} association may indicate "
                        f"{'mutual dependency or common underlying driver' if corr['correlation'] > 0 else 'compensatory or trade-off relationship'}.",
                        style='List Bullet'
                    )
            
            doc.add_paragraph()
            
            doc.add_paragraph(
                "Interpretation: These correlation patterns reveal systematic relationships within the dataset. "
                "Correlations above 0.7 warrant careful consideration in multivariate modeling due to potential "
                "multicollinearity. Positive correlations indicate variables that increase together, suggesting "
                "shared drivers or direct causal links. Negative correlations reveal inverse relationships that "
                "may represent trade-offs, constraints, or regulatory mechanisms."
            )
            
            doc.add_paragraph()
        
        # 4.5 Visual Analysis with Comprehensive Interpretations
        doc.add_heading('4.5 Visual Analysis and Graphical Interpretations', level=2)
        
        if charts and self.config.get('include_charts', True):
            for idx, chart in enumerate(charts, 1):
                if os.path.exists(chart['path']):
                    # Add chart
                    doc.add_paragraph(f"Figure {idx}: {chart['title']}", style='Caption')
                    doc.add_picture(chart['path'], width=Inches(6.0))
                    doc.add_paragraph()
                    
                    # Get interpretation from chart metadata
                    interpretation = chart.get('interpretation', '')
                    chart_title_lower = chart['title'].lower()
                    
                    # Add detailed observation, trend insight, and interpretation
                    if 'distribution' in chart_title_lower:
                        doc.add_paragraph(
                            f"Observation: {interpretation}"
                        )
                        doc.add_paragraph(
                            "Trend Insight: The distribution histogram reveals the frequency and spread of values "
                            "within the variable. The mean (red dashed line) indicates the arithmetic average, while "
                            "the median (green dashed line) shows the middle point of the data. When mean and median "
                            "diverge significantly, it suggests skewness—rightward skew if mean > median (influenced "
                            "by high outliers), leftward if mean < median (influenced by low outliers)."
                        )
                        doc.add_paragraph(
                            "Interpretation: The shape of the distribution provides critical insights into data behavior. "
                            "A bell-shaped curve suggests normal distribution suitable for parametric statistical tests. "
                            "Multi-modal patterns indicate distinct subgroups within the data. Heavy tails or extreme "
                            "values signal potential outliers requiring investigation for data quality or exceptional cases."
                        )
                    
                    elif 'correlation' in chart_title_lower or 'heatmap' in chart_title_lower:
                        doc.add_paragraph(
                            f"Observation: {interpretation}"
                        )
                        doc.add_paragraph(
                            "Trend Insight: The correlation matrix presents the strength and direction of linear "
                            "relationships between all numerical variable pairs. Color intensity indicates correlation "
                            "magnitude: darker colors (closer to ±1) represent strong correlations, while lighter "
                            "colors (closer to 0) indicate weak or no linear relationship. Warm colors typically "
                            "denote positive correlations (variables increase together), while cool colors indicate "
                            "negative correlations (inverse relationships)."
                        )
                        doc.add_paragraph(
                            "Interpretation: Strong correlations (|r| > 0.7) reveal variables that move in lock-step, "
                            "which is critical for understanding system dynamics and avoiding multicollinearity in "
                            "predictive models. Moderate correlations (0.4 < |r| < 0.7) suggest meaningful but not "
                            "dominant relationships. Weak correlations may still be statistically significant with "
                            "large sample sizes but have limited practical impact. Identifying correlation clusters "
                            "helps group related variables and understand underlying data structures."
                        )
                    
                    elif 'category' in chart_title_lower or 'bar' in chart_title_lower:
                        doc.add_paragraph(
                            f"Observation: {interpretation}"
                        )
                        doc.add_paragraph(
                            "Trend Insight: The categorical distribution displays frequency counts and percentage "
                            "composition across distinct categories. Bar height represents absolute counts, while "
                            "percentage labels provide relative proportions. This visualization immediately reveals "
                            "dominant categories, balanced vs. imbalanced distributions, and minority groups."
                        )
                        doc.add_paragraph(
                            "Interpretation: Highly skewed distributions (one category >50%) indicate concentration "
                            "that may affect representativeness and generalizability. Balanced distributions enable "
                            "robust comparative analysis across categories. Very small categories (<5%) may lack "
                            "statistical power for subgroup analysis and could be consolidated. Understanding category "
                            "distribution is essential for stratified sampling, quota setting, and ensuring diversity "
                            "in segmentation strategies."
                        )
                    
                    elif 'box plot' in chart_title_lower or 'boxplot' in chart_title_lower:
                        doc.add_paragraph(
                            f"Observation: {interpretation}"
                        )
                        doc.add_paragraph(
                            "Trend Insight: Box plots provide a five-number summary (minimum, Q1, median, Q3, maximum) "
                            "with outlier identification. The box represents the interquartile range (IQR) containing "
                            "the middle 50% of data. The line inside the box marks the median (50th percentile). "
                            "Whiskers extend to the furthest points within 1.5×IQR from quartiles. Points beyond "
                            "whiskers are potential outliers requiring investigation."
                        )
                        doc.add_paragraph(
                            "Interpretation: Box plots excel at revealing distribution shape, central tendency, spread, "
                            "and outliers simultaneously. Symmetric boxes suggest normal distributions; asymmetric boxes "
                            "indicate skewness. Long whiskers show high variability; short whiskers suggest concentrated "
                            "data. Outliers may represent data errors, exceptional cases, or population extremes—each "
                            "requiring different handling strategies. Comparing box plots across categories reveals "
                            "distributional differences critical for segmentation analysis."
                        )
                    
                    elif 'scatter' in chart_title_lower:
                        doc.add_paragraph(
                            f"Observation: {interpretation}"
                        )
                        doc.add_paragraph(
                            "Trend Insight: Scatter plots illustrate bivariate relationships with each point representing "
                            "an individual observation's values on two variables. The overall pattern reveals relationship "
                            "type: linear patterns suggest consistent proportional relationships, curved patterns indicate "
                            "non-linear associations, and random scatter suggests independence. The trend line (if present) "
                            "shows the best-fit linear approximation."
                        )
                        doc.add_paragraph(
                            "Interpretation: Tight clustering around the trend line indicates strong predictive "
                            "relationships where one variable reliably forecasts the other. Wide dispersion suggests "
                            "high residual variance and limited predictability. The slope direction shows relationship "
                            "type: positive slopes mean variables increase together, negative slopes indicate inverse "
                            "relationships. Identifying outliers—points far from the trend line—highlights unusual "
                            "cases that may represent errors, innovations, or special populations requiring separate analysis."
                        )
                    
                    else:
                        # Generic interpretation
                        doc.add_paragraph(f"Interpretation: {interpretation}")
                    
                    doc.add_paragraph()  # Spacing
        
        # 5. KEY INSIGHTS
        doc.add_page_break()
        doc.add_heading('5. Key Insights', level=1)
        
        if narrative and 'insights_interpretation' in narrative:
            doc.add_paragraph(narrative['insights_interpretation'])
        else:
            # Generate insights from data
            self._add_key_insights(doc, data_summary, numeric_summary, categorical_summary, correlations)
        
        # 6. RECOMMENDATIONS
        doc.add_page_break()
        doc.add_heading('6. Recommendations', level=1)
        
        if narrative and 'recommendations' in narrative:
            doc.add_paragraph(narrative['recommendations'])
        else:
            # Generate evidence-based recommendations
            self._add_recommendations(doc, data_summary, quality_score, correlations)
        
        # 7. CONCLUSION
        doc.add_page_break()
        doc.add_heading('7. Conclusion', level=1)
        
        self._add_comprehensive_conclusion(doc, data_summary, numeric_summary, categorical_summary, 
                                          correlations, quality_score, missing_pct)
        
        # 8. APPENDIX
        doc.add_page_break()
        doc.add_heading('8. Appendix — Summary Metrics', level=1)
        
        # Comprehensive appendix table
        appendix_table = doc.add_table(rows=8 + len(correlations[:3]), cols=2)
        appendix_table.style = 'Light Grid Accent 1'
        
        appendix_table.rows[0].cells[0].text = 'Metric'
        appendix_table.rows[0].cells[1].text = 'Value'
        for cell in appendix_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        appendix_table.rows[1].cells[0].text = 'Total Records'
        appendix_table.rows[1].cells[1].text = f"{total_rows:,}"
        appendix_table.rows[2].cells[0].text = 'Total Variables'
        appendix_table.rows[2].cells[1].text = f"{total_cols}"
        appendix_table.rows[3].cells[0].text = 'Numerical Variables'
        appendix_table.rows[3].cells[1].text = f"{len(numeric_cols)}"
        appendix_table.rows[4].cells[0].text = 'Categorical Variables'
        appendix_table.rows[4].cells[1].text = f"{len(categorical_cols)}"
        appendix_table.rows[5].cells[0].text = 'Data Quality Score'
        appendix_table.rows[5].cells[1].text = f"{quality_score:.1f}/100"
        appendix_table.rows[6].cells[0].text = 'Missing Data'
        appendix_table.rows[6].cells[1].text = f"{missing_pct:.2f}%"
        appendix_table.rows[7].cells[0].text = 'Analysis Date'
        appendix_table.rows[7].cells[1].text = datetime.now().strftime('%Y-%m-%d')
        
        # Add top correlations
        for idx, corr in enumerate(correlations[:3], start=8):
            appendix_table.rows[idx].cells[0].text = f"Correlation: {corr['var1']} ↔ {corr['var2']}"
            appendix_table.rows[idx].cells[1].text = f"{corr['correlation']:.3f}"
        
        # Save document
        doc.save(filepath)
        return filepath
    
    def _add_executive_summary(self, doc, data_summary, numeric_summary, categorical_summary, correlations):
        """Generate comprehensive executive summary from data"""
        total_rows = data_summary.get('total_rows', 0)
        total_cols = data_summary.get('total_columns', 0)
        quality_score = data_summary.get('data_quality_score', 0)
        
        doc.add_paragraph(
            f"This comprehensive analysis examines {total_rows:,} records across {total_cols} variables, "
            f"revealing key patterns, relationships, and actionable insights:"
        )
        
        summary_points = []
        
        # Data quality
        summary_points.append(
            f"Data Integrity: Quality score of {quality_score:.1f}/100 indicates "
            f"{'excellent' if quality_score >= 90 else 'good' if quality_score >= 75 else 'acceptable'} "
            f"data reliability, providing a solid analytical foundation."
        )
        
        # Numerical insights
        if numeric_summary:
            first_num = list(numeric_summary.items())[0]
            var_name, stats = first_num
            summary_points.append(
                f"Primary Metric ({var_name}): Average value of {stats.get('mean', 0):,.2f} "
                f"with median at {stats.get('median', 0):,.2f}, ranging from {stats.get('min', 0):,.2f} "
                f"to {stats.get('max', 0):,.2f}."
            )
        
        # Correlation insights
        if correlations:
            strong_corrs = [c for c in correlations if abs(c['correlation']) > 0.5]
            if strong_corrs:
                top_corr = strong_corrs[0]
                summary_points.append(
                    f"Strongest Relationship: {top_corr['var1']} and {top_corr['var2']} show "
                    f"{'strong positive' if top_corr['correlation'] > 0 else 'strong negative'} "
                    f"correlation (r = {top_corr['correlation']:.3f}), indicating significant interdependence."
                )
        
        # Categorical insights
        if categorical_summary:
            first_cat = list(categorical_summary.items())[0]
            var_name, stats = first_cat
            summary_points.append(
                f"Categorical Distribution ({var_name}): {stats.get('unique_values', 0)} distinct "
                f"categories with {stats.get('top_value_percentage', 0):.1f}% concentration in dominant category."
            )
        
        # Variable composition
        num_count = len(data_summary.get('numeric_column_names', []))
        cat_count = len(data_summary.get('categorical_column_names', []))
        summary_points.append(
            f"Variable Composition: {num_count} numerical metrics and {cat_count} categorical dimensions "
            f"enable comprehensive quantitative and qualitative analysis."
        )
        
        for point in summary_points:
            doc.add_paragraph(point, style='List Bullet')
    
    def _add_key_insights(self, doc, data_summary, numeric_summary, categorical_summary, correlations):
        """Generate key insights section"""
        insights = []
        
        quality_score = data_summary.get('data_quality_score', 0)
        total_rows = data_summary.get('total_rows', 0)
        
        insights.append(
            f"Data Quality: Quality score of {quality_score:.1f}/100 with {total_rows:,} records "
            f"demonstrates {'excellent' if quality_score >= 90 else 'solid'} dataset integrity suitable "
            f"for strategic decision-making."
        )
        
        if correlations:
            strong_count = len([c for c in correlations if abs(c['correlation']) > 0.5])
            insights.append(
                f"Variable Relationships: Identified {strong_count} strong correlations (|r| > 0.5) revealing "
                f"significant interdependencies that inform system understanding and predictive modeling."
            )
        
        if numeric_summary:
            insights.append(
                f"Numerical Patterns: Analysis of {len(numeric_summary)} numerical variables reveals "
                f"distinct statistical profiles with varying degrees of dispersion and distribution characteristics."
            )
        
        if categorical_summary:
            insights.append(
                f"Categorical Diversity: {len(categorical_summary)} categorical variables provide "
                f"multidimensional segmentation capabilities for targeted analysis and strategic planning."
            )
        
        doc.add_paragraph("Key analytical insights from this comprehensive examination:")
        for insight in insights:
            doc.add_paragraph(insight, style='List Bullet')
    
    def _add_recommendations(self, doc, data_summary, quality_score, correlations):
        """Generate evidence-based recommendations"""
        doc.add_paragraph(
            "Based on the comprehensive analysis, the following evidence-based recommendations are proposed:"
        )
        
        recommendations = []
        
        if quality_score < 90:
            recommendations.append({
                'focus': 'Data Quality Enhancement',
                'recommendation': 'Implement systematic data validation protocols and missing value reduction strategies.',
                'outcome': f'Improve data quality score from {quality_score:.1f} to >90 within 6 months.'
            })
        
        if correlations:
            strong_corrs = [c for c in correlations if abs(c['correlation']) > 0.7]
            if strong_corrs:
                recommendations.append({
                    'focus': 'Relationship Monitoring',
                    'recommendation': 'Establish tracking mechanisms for strongly correlated variables to detect systemic changes.',
                    'outcome': 'Enable early warning of structural shifts affecting multiple dimensions simultaneously.'
                })
        
        recommendations.append({
            'focus': 'Continuous Analytics',
            'recommendation': 'Implement periodic re-analysis (quarterly) to track trends and validate insights over time.',
            'outcome': 'Maintain current understanding and adapt strategies to evolving patterns.'
        })
        
        recommendations.append({
            'focus': 'Stakeholder Engagement',
            'recommendation': 'Conduct findings validation sessions with subject matter experts and operational leaders.',
            'outcome': 'Ensure analytical insights align with operational realities and strategic objectives.'
        })
        
        if recommendations:
            # Create recommendations table
            rec_table = doc.add_table(rows=len(recommendations) + 1, cols=3)
            rec_table.style = 'Light Grid Accent 1'
            
            rec_table.rows[0].cells[0].text = 'Focus Area'
            rec_table.rows[0].cells[1].text = 'Recommendation'
            rec_table.rows[0].cells[2].text = 'Expected Outcome'
            
            for cell in rec_table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            for idx, rec in enumerate(recommendations, start=1):
                rec_table.rows[idx].cells[0].text = rec['focus']
                rec_table.rows[idx].cells[1].text = rec['recommendation']
                rec_table.rows[idx].cells[2].text = rec['outcome']
    
    def _add_comprehensive_conclusion(self, doc, data_summary, numeric_summary, categorical_summary, 
                                     correlations, quality_score, missing_pct):
        """Generate comprehensive conclusion matching HR example standard"""
        total_rows = data_summary.get('total_rows', 0)
        total_cols = data_summary.get('total_columns', 0)
        
        doc.add_paragraph(
            f"This comprehensive analytical examination of {total_rows:,} records across {total_cols} "
            f"variables demonstrates a {'robust' if quality_score >= 85 else 'solid'} dataset with "
            f"well-defined patterns and relationships. The analysis reveals:"
        )
        
        conclusion_points = []
        
        # Data quality assessment
        conclusion_points.append(
            f"Data Foundation: Quality score of {quality_score:.1f}/100 with {missing_pct:.2f}% missing data "
            f"establishes {'excellent' if quality_score >= 90 else 'good'} analytical reliability."
        )
        
        # Statistical insights
        if numeric_summary:
            conclusion_points.append(
                f"Numerical Profile: {len(numeric_summary)} quantitative variables exhibit distinct statistical "
                f"characteristics with varying levels of central tendency, dispersion, and distributional properties."
            )
        
        # Categorical insights
        if categorical_summary:
            conclusion_points.append(
                f"Categorical Landscape: {len(categorical_summary)} qualitative dimensions provide rich "
                f"segmentation capabilities for targeted interventions and strategic planning."
            )
        
        # Correlation insights
        if correlations:
            strong_count = len([c for c in correlations if abs(c['correlation']) > 0.5])
            conclusion_points.append(
                f"Systemic Relationships: {strong_count} strong correlations reveal interconnected dynamics "
                f"requiring holistic rather than isolated intervention approaches."
            )
        
        for point in conclusion_points:
            doc.add_paragraph(point, style='List Bullet')
        
        doc.add_paragraph()
        doc.add_paragraph(
            "To maximize strategic value from these insights:"
        )
        
        next_steps = [
            "Maintain data quality through continuous validation and enhancement protocols.",
            "Monitor key correlations for structural changes indicating systemic shifts.",
            "Conduct periodic re-analysis to track trends and validate strategic assumptions.",
            "Integrate findings into operational planning and performance management frameworks."
        ]
        
        for step in next_steps:
            doc.add_paragraph(step, style='List Bullet')
        
        doc.add_paragraph()
        doc.add_paragraph(
            "Overall, TK Reports concludes the dataset provides a reliable foundation for evidence-based "
            "decision-making, with clear patterns supporting strategic initiatives and operational optimization."
        )
    
    def _add_cover_page_docx(self, doc, metadata):
        """Create professional cover page"""
        # Title
        title = doc.add_heading(metadata.get('report_title', 'TK Reports — Comprehensive Data Insight Report'), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph()
        doc.add_paragraph()
        
        prepared_for = doc.add_paragraph()
        prepared_for.add_run('Prepared for: ').bold = True
        prepared_for.add_run(metadata.get('prepared_for', 'Organization'))
        prepared_for.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        prepared_by = doc.add_paragraph()
        prepared_by.add_run('Prepared by: ').bold = True
        prepared_by.add_run(metadata.get('prepared_by', 'TK Reports (AI-Assisted Insight Reporting System)'))
        prepared_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph()
        date_para.add_run('Date: ').bold = True
        date_para.add_run(datetime.now().strftime('%B %d, %Y'))
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        region_para = doc.add_paragraph()
        region_para.add_run('Region: ').bold = True
        region_para.add_run('en_CA')
        region_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def create_pdf_report(self, filepath, data_summary, narrative, charts, metadata):
        """Create comprehensive PDF report with robust error handling"""
        try:
            pdf = FPDF('P', 'mm', 'A4')  # Portrait, millimeters, A4 size
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_margins(left=20, top=20, right=20)
            
            # Extract key metrics
            total_rows = data_summary.get('total_rows', 0)
            total_cols = data_summary.get('total_columns', 0)
            quality_score = data_summary.get('data_quality_score', 0)
            missing_pct = data_summary.get('missing_percentage', 0)
            numeric_summary = data_summary.get('numeric_summary', {})
            categorical_summary = data_summary.get('categorical_summary', {})
            correlations = data_summary.get('correlations', [])
            
            # Cover Page
            pdf.add_page()
            pdf.set_font('Helvetica', 'B', 18)
            title_text = self._clean_text_for_pdf(metadata.get('report_title', 'TK Reports - Data Insight Report'))
            if len(title_text) > 60:
                title_text = title_text[:57] + '...'
            
            # Center title with proper width
            pdf.cell(0, 15, title_text, align='C', ln=True)
            pdf.ln(10)
            
            pdf.set_font('Helvetica', 'B', 11)
            pdf.cell(0, 8, f"Prepared for: {metadata.get('prepared_for', 'Organization')[:40]}", align='C', ln=True)
            pdf.cell(0, 8, f"Prepared by: {metadata.get('prepared_by', 'TK Reports')[:40]}", align='C', ln=True)
            
            pdf.set_font('Helvetica', '', 10)
            pdf.cell(0, 8, f"Date: {datetime.now().strftime('%B %d, %Y')}", align='C', ln=True)
            pdf.cell(0, 8, "Region: en_CA", align='C', ln=True)
            
            # Executive Summary
            pdf.add_page()
            pdf.set_font('Helvetica', 'B', 13)
            pdf.cell(0, 10, '1. Executive Summary', ln=True)
            pdf.set_font('Helvetica', '', 9)
            
            if narrative and 'executive_summary' in narrative:
                summary_text = self._clean_text_for_pdf(narrative['executive_summary'][:1200])
                pdf.multi_cell(0, 5, summary_text)
            else:
                pdf.multi_cell(0, 5, 
                    f"This analysis examines {total_rows:,} records across {total_cols} variables. "
                    f"Data quality score: {quality_score:.1f}/100."
                )
            
            # Introduction
            pdf.add_page()
            pdf.set_font('Helvetica', 'B', 13)
            pdf.cell(0, 10, '2. Introduction', ln=True)
            
            # Add Research Objective if provided
            if metadata.get('objective'):
                pdf.set_font('Helvetica', 'B', 10)
                pdf.cell(0, 6, '2.1 Research Objective and Purpose', ln=True)
                pdf.set_font('Helvetica', '', 9)
                objective_text = self._clean_text_for_pdf(metadata['objective'][:500])
                pdf.multi_cell(0, 5, objective_text)
                pdf.ln(2)
                
                pdf.set_font('Helvetica', 'B', 10)
                pdf.cell(0, 6, '2.2 Dataset Overview', ln=True)
                pdf.set_font('Helvetica', '', 9)
            
            pdf.multi_cell(0, 5,
                f"Dataset: {total_rows:,} records with {total_cols} variables for comprehensive analysis."
            )
            
            # Methodology
            pdf.add_page()
            pdf.set_font('Helvetica', 'B', 13)
            pdf.cell(0, 10, '3. Methodology', ln=True)
            
            pdf.set_font('Helvetica', 'B', 10)
            pdf.cell(0, 6, '3.1 Data Source and Scope', ln=True)
            pdf.set_font('Helvetica', '', 9)
            pdf.multi_cell(0, 5,
                f"{total_rows:,} records analyzed with {missing_pct:.2f}% missing values."
            )
            
            pdf.ln(2)
            pdf.set_font('Helvetica', 'B', 10)
            pdf.cell(0, 6, '3.2 Analytical Process', ln=True)
            pdf.set_font('Helvetica', '', 9)
            pdf.multi_cell(0, 5,
                "Multi-stage approach: Statistics, Distribution, Correlation, Visual Analytics, AI Interpretation."
            )
            
            # Detailed Findings
            pdf.add_page()
            pdf.set_font('Helvetica', 'B', 13)
            pdf.cell(0, 10, '4. Detailed Findings and Interpretations', ln=True)
            
            if narrative and 'key_findings' in narrative:
                pdf.set_font('Helvetica', '', 9)
                findings_text = self._clean_text_for_pdf(narrative['key_findings'][:1200])
                pdf.multi_cell(0, 5, findings_text)
            else:
                pdf.set_font('Helvetica', '', 9)
                pdf.multi_cell(0, 5,
                    f"{total_rows:,} records with {total_cols} variables. Quality: {quality_score:.1f}/100."
                )
            
            # Numerical Statistics Summary (SIMPLIFIED to avoid errors)
            if numeric_summary:
                pdf.ln(3)
                pdf.set_font('Helvetica', 'B', 10)
                pdf.cell(0, 6, '4.1 Key Numerical Variables', ln=True)
                pdf.set_font('Helvetica', '', 8)
                
                for var_name, stats in list(numeric_summary.items())[:5]:
                    try:
                        mean_val = stats.get('mean', 0)
                        median_val = stats.get('median', 0)
                        
                        # Very short display name
                        display_name = var_name[:20] if len(var_name) <= 20 else var_name[:17] + '...'
                        display_name = self._clean_text_for_pdf(display_name)
                        
                        # Simple format
                        pdf.cell(0, 4, f"- {display_name}: Mean={mean_val:.2f}, Median={median_val:.2f}", ln=True)
                    except Exception:
                        continue  # Skip problematic variables
            
            # Correlations (SIMPLIFIED)
            if correlations:
                pdf.ln(3)
                pdf.set_font('Helvetica', 'B', 10)
                pdf.cell(0, 6, '4.2 Key Correlations', ln=True)
                pdf.set_font('Helvetica', '', 8)
                
                for corr in correlations[:5]:
                    try:
                        var1 = corr['var1'][:15] if len(corr['var1']) <= 15 else corr['var1'][:12] + '...'
                        var2 = corr['var2'][:15] if len(corr['var2']) <= 15 else corr['var2'][:12] + '...'
                        var1 = self._clean_text_for_pdf(var1)
                        var2 = self._clean_text_for_pdf(var2)
                        
                        pdf.cell(0, 4, f"- {var1} <-> {var2}: r={corr['correlation']:.3f}", ln=True)
                    except Exception:
                        continue
            
            # Charts - With robust error handling
            if charts and self.config.get('include_charts', True):
                pdf.add_page()
                pdf.set_font('Helvetica', 'B', 13)
                pdf.cell(0, 10, '4.3 Visual Analysis', ln=True)
                
                charts_added = 0
                for idx, chart in enumerate(charts, 1):
                    if os.path.exists(chart['path']):
                        try:
                            # Add chart title
                            pdf.ln(3)
                            pdf.set_font('Helvetica', 'B', 10)
                            chart_title = chart['title'][:50] if len(chart['title']) <= 50 else chart['title'][:47] + '...'
                            pdf.cell(0, 6, f"Figure {idx}: {self._clean_text_for_pdf(chart_title)}", ln=True)
                            
                            # Calculate available width (A4 width is 210mm, minus margins)
                            page_width = 210  # A4 width in mm
                            margin_total = 40  # 20mm left + 20mm right
                            img_width = page_width - margin_total  # 170mm
                            
                            # Add the image
                            current_y = pdf.get_y()
                            pdf.image(chart['path'], x=20, y=current_y, w=img_width)
                            
                            # Move cursor below image (estimate height as 100mm)
                            pdf.set_y(current_y + 100)
                            
                            # Add interpretation if available
                            if chart.get('interpretation'):
                                pdf.set_font('Helvetica', '', 8)
                                interp_text = self._clean_text_for_pdf(chart['interpretation'][:200])
                                pdf.multi_cell(0, 4, f"Interpretation: {interp_text}...")
                            
                            pdf.ln(3)
                            charts_added += 1
                            
                            # Add page break after every 2 charts
                            if charts_added % 2 == 0 and idx < len(charts):
                                pdf.add_page()
                                
                        except Exception as chart_error:
                            # Silently continue if chart fails
                            print(f"Warning: Could not add chart {idx} to PDF: {str(chart_error)}")
                            pdf.set_font('Helvetica', 'I', 8)
                            pdf.cell(0, 5, f"[Chart {idx} visualization available in DOCX report]", ln=True)
                            pdf.ln(2)
                
                # If no charts were added, show message
                if charts_added == 0:
                    pdf.set_font('Helvetica', '', 9)
                    pdf.multi_cell(0, 5, 
                        f"{len(charts)} visualizations generated. Please see DOCX report for full charts."
                    )
            
            # Key Insights
            pdf.add_page()
            pdf.set_font('Helvetica', 'B', 13)
            pdf.cell(0, 10, '5. Key Insights', ln=True)
            pdf.set_font('Helvetica', '', 9)
            
            if narrative and 'insights_interpretation' in narrative:
                insights_text = self._clean_text_for_pdf(narrative['insights_interpretation'][:1200])
                pdf.multi_cell(0, 5, insights_text)
            else:
                pdf.multi_cell(0, 5,
                    f"Analysis reveals patterns across {total_rows:,} records. Quality score: {quality_score:.1f}/100."
                )
            
            # Recommendations
            pdf.add_page()
            pdf.set_font('Helvetica', 'B', 13)
            pdf.cell(0, 10, '6. Recommendations', ln=True)
            pdf.set_font('Helvetica', '', 9)
            
            if narrative and 'recommendations' in narrative:
                rec_text = self._clean_text_for_pdf(narrative['recommendations'][:1200])
                pdf.multi_cell(0, 5, rec_text)
            else:
                pdf.multi_cell(0, 5,
                    "Implement data quality protocols and periodic re-analysis."
                )
            
            # Conclusion
            pdf.add_page()
            pdf.set_font('Helvetica', 'B', 13)
            pdf.cell(0, 10, '7. Conclusion', ln=True)
            pdf.set_font('Helvetica', '', 9)
            pdf.multi_cell(0, 5,
                f"Analysis of {total_rows:,} records demonstrates "
                f"{'robust' if quality_score >= 85 else 'solid'} dataset with reliable patterns."
            )
            
            # Appendix
            pdf.add_page()
            pdf.set_font('Helvetica', 'B', 13)
            pdf.cell(0, 10, '8. Appendix - Summary Metrics', ln=True)
            pdf.set_font('Helvetica', '', 9)
            
            pdf.cell(0, 5, f"Total Records: {total_rows:,}", ln=True)
            pdf.cell(0, 5, f"Total Variables: {total_cols}", ln=True)
            pdf.cell(0, 5, f"Data Quality Score: {quality_score:.1f}/100", ln=True)
            pdf.cell(0, 5, f"Missing Data: {missing_pct:.2f}%", ln=True)
            pdf.cell(0, 5, f"Analysis Date: {datetime.now().strftime('%Y-%m-%d')}", ln=True)
            
            # Save PDF
            pdf.output(filepath)
            print(f"PDF report saved successfully to: {filepath}")
            return filepath
            
        except Exception as e:
            print(f"Error creating PDF report: {str(e)}")
            import traceback
            traceback.print_exc()
            return None
            
            # Extract key metrics
            total_rows = data_summary.get('total_rows', 0)
            total_cols = data_summary.get('total_columns', 0)
            quality_score = data_summary.get('data_quality_score', 0)
            missing_pct = data_summary.get('missing_percentage', 0)
            numeric_summary = data_summary.get('numeric_summary', {})
            categorical_summary = data_summary.get('categorical_summary', {})
            correlations = data_summary.get('correlations', [])
            
            # Cover Page
            pdf.add_page()
            pdf.set_font('Arial', 'B', 20)
            title_text = self._clean_text_for_pdf(metadata.get('report_title', 'TK Reports - Comprehensive Data Insight Report'))
            # Truncate if too long
            if len(title_text) > 80:
                title_text = title_text[:77] + '...'
            pdf.multi_cell(0, 10, title_text, 0, 'C')
            pdf.ln(10)
            
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(0, 8, self._clean_text_for_pdf(f"Prepared for: {metadata.get('prepared_for', 'Organization')[:50]}"), 0, 1, 'C')
            pdf.cell(0, 8, self._clean_text_for_pdf(f"Prepared by: {metadata.get('prepared_by', 'TK Reports')[:50]}"), 0, 1, 'C')
            
            pdf.set_font('Arial', '', 11)
            pdf.cell(0, 8, f"Date: {datetime.now().strftime('%B %d, %Y')}", 0, 1, 'C')
            pdf.cell(0, 8, "Region: en_CA", 0, 1, 'C')
            
            # Executive Summary
            pdf.add_page()
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, '1. Executive Summary', 0, 1)
            pdf.set_font('Arial', '', 10)
            
            if narrative and 'executive_summary' in narrative:
                summary_text = self._clean_text_for_pdf(narrative['executive_summary'][:1500])
                pdf.multi_cell(0, 5, summary_text)
            else:
                pdf.multi_cell(0, 5, self._clean_text_for_pdf(
                    f"This comprehensive analysis examines {total_rows:,} records across {total_cols} variables, "
                    f"revealing key patterns and relationships. Data quality score of {quality_score:.1f}/100 "
                    f"demonstrates {'excellent' if quality_score >= 90 else 'good'} integrity."
                ))
            
            # Introduction
            pdf.add_page()
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, '2. Introduction', 0, 1)
            pdf.set_font('Arial', '', 10)
            pdf.multi_cell(0, 5, self._clean_text_for_pdf(
                f"This comprehensive analytical report examines a dataset containing {total_rows:,} records "
                f"across {total_cols} variables. The dataset provides a multifaceted view suitable for "
                f"comprehensive analysis and actionable insights."
            ))
            
            # Methodology
            pdf.add_page()
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, '3. Methodology', 0, 1)
            
            pdf.set_font('Arial', 'B', 11)
            pdf.cell(0, 7, '3.1 Data Source and Scope', 0, 1)
            pdf.set_font('Arial', '', 10)
            pdf.multi_cell(0, 5, self._clean_text_for_pdf(
                f"Source dataset comprises {total_rows:,} records with {total_cols} variables. "
                f"Analysis conducted on numerical and categorical dimensions with {missing_pct:.2f}% missing values."
            ))
            
            pdf.ln(3)
            pdf.set_font('Arial', 'B', 11)
            pdf.cell(0, 7, '3.2 Analytical Process', 0, 1)
            pdf.set_font('Arial', '', 10)
            pdf.multi_cell(0, 5, self._clean_text_for_pdf(
                "Multi-stage approach: Descriptive Statistics, Distribution Analysis, Advanced Statistical Metrics, "
                "Correlation Analysis, Categorical Profiling, Visual Analytics, and AI-Powered Interpretation."
            ))
            
            # Detailed Findings
            pdf.add_page()
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, '4. Detailed Findings and Interpretations', 0, 1)
            
            if narrative and 'key_findings' in narrative:
                pdf.set_font('Arial', '', 10)
                findings_text = self._clean_text_for_pdf(narrative['key_findings'][:1500])
                pdf.multi_cell(0, 5, findings_text)
            else:
                pdf.set_font('Arial', '', 10)
                pdf.multi_cell(0, 5, self._clean_text_for_pdf(
                    f"Dataset encompasses {total_rows:,} records representing the analytical universe, with {total_cols} "
                    f"variables providing multidimensional perspectives. Quality score of {quality_score:.1f}/100 "
                    f"demonstrates robust analytical foundation."
                ))
            
            # Numerical Statistics Summary
            if numeric_summary:
                pdf.ln(5)
                pdf.set_font('Arial', 'B', 11)
                pdf.cell(0, 7, '4.1 Key Numerical Variables', 0, 1)
                pdf.set_font('Arial', '', 9)
                
                for var_name, stats in list(numeric_summary.items())[:5]:
                    mean_val = stats.get('mean', 0)
                    median_val = stats.get('median', 0)
                    std_val = stats.get('std', 0)
                    
                    # Truncate variable name if too long
                    display_name = var_name[:40] + '...' if len(var_name) > 40 else var_name
                    
                    text = self._clean_text_for_pdf(
                        f"- {display_name}: Mean={mean_val:,.2f}, Median={median_val:,.2f}, StdDev={std_val:,.2f}"
                    )
                    pdf.multi_cell(0, 4, text)
            
            # Correlations
            if correlations:
                pdf.ln(5)
                pdf.set_font('Arial', 'B', 11)
                pdf.cell(0, 7, '4.2 Key Correlations', 0, 1)
                pdf.set_font('Arial', '', 9)
                
                for corr in correlations[:5]:
                    var1 = corr['var1'][:25] + '...' if len(corr['var1']) > 25 else corr['var1']
                    var2 = corr['var2'][:25] + '...' if len(corr['var2']) > 25 else corr['var2']
                    
                    text = self._clean_text_for_pdf(
                        f"- {var1} <-> {var2}: r={corr['correlation']:.3f} ({corr.get('strength', 'N/A')})"
                    )
                    pdf.multi_cell(0, 4, text)
            
            # Charts - with better error handling
            if charts and self.config.get('include_charts', True):
                pdf.add_page()
                pdf.set_font('Arial', 'B', 14)
                pdf.cell(0, 10, '4.3 Visual Analysis', 0, 1)
                
                for idx, chart in enumerate(charts[:3], 1):  # Limit to 3 charts for PDF
                    if os.path.exists(chart['path']):
                        try:
                            # Check if new page needed
                            if pdf.get_y() > 200:
                                pdf.add_page()
                            
                            pdf.ln(3)
                            pdf.set_font('Arial', 'B', 10)
                            chart_title = chart['title'][:60] + '...' if len(chart['title']) > 60 else chart['title']
                            pdf.cell(0, 7, self._clean_text_for_pdf(f"Figure {idx}: {chart_title}"), 0, 1)
                            
                            # Add image with error handling
                            try:
                                # Calculate available width (page width - margins)
                                available_width = pdf.w - 30  # Leave margins
                                pdf.image(chart['path'], x=15, w=available_width)
                                pdf.ln(3)
                            except Exception as img_error:
                                pdf.set_font('Arial', 'I', 9)
                                pdf.multi_cell(0, 5, f"[Chart image could not be embedded: {str(img_error)[:50]}]")
                            
                            # Add interpretation
                            if chart.get('interpretation'):
                                pdf.set_font('Arial', '', 9)
                                interp_text = self._clean_text_for_pdf(chart['interpretation'][:250])
                                pdf.multi_cell(0, 4, f"Interpretation: {interp_text}...")
                                pdf.ln(2)
                                
                        except Exception as chart_error:
                            pdf.set_font('Arial', 'I', 9)
                            pdf.multi_cell(0, 5, f"[Error rendering chart {idx}: {str(chart_error)[:50]}]")
                            pdf.ln(2)
            
            # Key Insights
            pdf.add_page()
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, '5. Key Insights', 0, 1)
            pdf.set_font('Arial', '', 10)
            
            if narrative and 'insights_interpretation' in narrative:
                insights_text = self._clean_text_for_pdf(narrative['insights_interpretation'][:1500])
                pdf.multi_cell(0, 5, insights_text)
            else:
                pdf.multi_cell(0, 5, self._clean_text_for_pdf(
                    f"Analysis reveals significant patterns across {total_rows:,} records. "
                    f"Data quality score of {quality_score:.1f}/100 provides reliable foundation for decision-making."
                ))
            
            # Recommendations
            pdf.add_page()
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, '6. Recommendations', 0, 1)
            pdf.set_font('Arial', '', 10)
            
            if narrative and 'recommendations' in narrative:
                rec_text = self._clean_text_for_pdf(narrative['recommendations'][:1500])
                pdf.multi_cell(0, 5, rec_text)
            else:
                pdf.multi_cell(0, 5, self._clean_text_for_pdf(
                    "Based on comprehensive analysis, implement data quality enhancement protocols, "
                    "establish monitoring for key correlations, and conduct periodic re-analysis."
                ))
            
            # Conclusion
            pdf.add_page()
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, '7. Conclusion', 0, 1)
            pdf.set_font('Arial', '', 10)
            pdf.multi_cell(0, 5, self._clean_text_for_pdf(
                f"This comprehensive analysis of {total_rows:,} records across {total_cols} variables demonstrates "
                f"a {'robust' if quality_score >= 85 else 'solid'} dataset with well-defined patterns. "
                f"The findings provide reliable foundation for evidence-based decision-making."
            ))
            
            # Appendix
            pdf.add_page()
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, '8. Appendix - Summary Metrics', 0, 1)
            pdf.set_font('Arial', '', 9)
            
            appendix_text = self._clean_text_for_pdf(
                f"Total Records: {total_rows:,}\n"
                f"Total Variables: {total_cols}\n"
                f"Data Quality Score: {quality_score:.1f}/100\n"
                f"Missing Data: {missing_pct:.2f}%\n"
                f"Analysis Date: {datetime.now().strftime('%Y-%m-%d')}"
            )
            pdf.multi_cell(0, 5, appendix_text)
            
            # Save PDF
            pdf.output(filepath)
            return filepath
            
        except Exception as e:
            print(f"Error creating PDF report: {str(e)}")
            import traceback
            traceback.print_exc()
            return None
    
    def _clean_text_for_pdf(self, text):
        """Clean text for PDF compatibility"""
        replacements = {
            '•': '-',
            '→': '->',
            '↔': '<->',
            ''': "'",
            ''': "'",
            '"': '"',
            '"': '"',
            '–': '-',
            '—': '-',
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        return text
